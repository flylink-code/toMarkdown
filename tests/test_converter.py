"""Tests for docx to markdown conversion."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from tomarkdown.converter import (
    ConversionError,
    convert_batch,
    convert_docx_to_md,
    convert_file,
)

SAMPLE_DOCX = Path(__file__).parent / "sample.docx"
SAMPLE_PDF = Path(__file__).parent / "sample.pdf"
SAMPLE_TITLE = "Hello Markdown"
SAMPLE_PDF_TEXT = "Hello PDF"


def _write_sample_pdf(path: Path) -> None:
    """Write a minimal text-based PDF for conversion tests."""
    stream = "BT /F1 24 Tf 100 700 Td (Hello PDF) Tj ET"
    objects = [
        "1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj",
        "2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj",
        (
            "3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R"
            "/Resources<</Font<</F1 4 0 R>>>>/Contents 5 0 R>>endobj"
        ),
        "4 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj",
        f"5 0 obj<</Length {len(stream)}>>stream\n{stream}\nendstream endobj",
    ]

    body = ["%PDF-1.4"]
    offsets = [0]
    for obj in objects:
        offsets.append(sum(len(part) + 1 for part in body))
        body.append(obj)

    xref_offset = sum(len(part) + 1 for part in body)
    body.append("xref")
    body.append(f"0 {len(objects) + 1}")
    body.append("0000000000 65535 f ")
    for offset in offsets[1:]:
        body.append(f"{offset:010d} 00000 n ")
    body.append("trailer<</Size 6/Root 1 0 R>>")
    body.append(f"startxref\n{xref_offset}")
    body.append("%%EOF")
    path.write_text("\n".join(body), encoding="latin-1")


@pytest.fixture(scope="module", autouse=True)
def create_sample_docx() -> None:
    """Create sample input files for tests if they do not exist."""
    if not SAMPLE_DOCX.exists():
        doc = Document()
        doc.add_heading(SAMPLE_TITLE, level=1)
        doc.add_paragraph("This is a test paragraph for conversion.")
        doc.add_paragraph("Second line with bold text.", style="Normal")
        doc.save(SAMPLE_DOCX)

    if not SAMPLE_PDF.exists():
        _write_sample_pdf(SAMPLE_PDF)


def test_convert_pdf_to_md(tmp_path: Path) -> None:
    """Given a text-based .pdf file, output a valid .md file."""
    pytest.importorskip("pdfminer")
    pytest.importorskip("pdfplumber")

    dst = tmp_path / "output.md"
    convert_docx_to_md(SAMPLE_PDF, dst)

    assert dst.exists()
    content = dst.read_text(encoding="utf-8")
    assert content.strip()
    assert SAMPLE_PDF_TEXT in content


def test_convert_docx_to_md(tmp_path: Path) -> None:
    """Given a .docx file, output a valid .md file."""
    dst = tmp_path / "output.md"

    convert_docx_to_md(SAMPLE_DOCX, dst)

    assert dst.exists()
    content = dst.read_text(encoding="utf-8")
    assert content.strip()
    assert SAMPLE_TITLE in content or "Hello" in content


def test_convert_missing_file(tmp_path: Path) -> None:
    """Raise FileNotFoundError when source does not exist."""
    with pytest.raises(FileNotFoundError):
        convert_docx_to_md(tmp_path / "missing.docx", tmp_path / "out.md")


def test_convert_unsupported_format(tmp_path: Path) -> None:
    """Raise ValueError for unsupported extensions."""
    src = tmp_path / "notes.txt"
    src.write_text("not a docx", encoding="utf-8")

    with pytest.raises(ValueError, match="Unsupported file format"):
        convert_docx_to_md(src, tmp_path / "out.md")


def test_convert_creates_output_directory(tmp_path: Path) -> None:
    """Create nested output directories as needed."""
    dst = tmp_path / "nested" / "dir" / "output.md"

    convert_docx_to_md(SAMPLE_DOCX, dst)

    assert dst.exists()


def test_convert_file_returns_bool(tmp_path: Path) -> None:
    """convert_file returns True on success."""
    dst = tmp_path / "output.md"
    assert convert_file(SAMPLE_DOCX, dst) is True
    assert dst.exists()


def test_convert_file_returns_false_on_missing(tmp_path: Path) -> None:
    """convert_file returns False when source is missing."""
    assert convert_file(tmp_path / "missing.docx", tmp_path / "out.md") is False


def test_convert_batch_skips_nonempty_output(tmp_path: Path) -> None:
    """Skip conversion when a non-empty output file already exists."""
    out_dir = tmp_path / "output"
    out_dir.mkdir()
    existing = out_dir / "sample.md"
    existing.write_text("existing content", encoding="utf-8")

    result = convert_batch([SAMPLE_DOCX], out_dir, overwrite=False)

    assert result["skipped_count"] == 1
    assert result["success_count"] == 0
    assert existing.read_text(encoding="utf-8") == "existing content"


def test_convert_batch_reconverts_empty_output(tmp_path: Path) -> None:
    """Re-convert when an existing output file is empty."""
    out_dir = tmp_path / "output"
    out_dir.mkdir()
    existing = out_dir / "sample.md"
    existing.write_text("", encoding="utf-8")

    result = convert_batch([SAMPLE_DOCX], out_dir, overwrite=False)

    assert result["skipped_count"] == 0
    assert result["success_count"] == 1
    assert existing.read_text(encoding="utf-8").strip()


def test_convert_batch_with_callback(tmp_path: Path) -> None:
    """convert_batch reports progress via callback."""
    out_dir = tmp_path / "output"
    progress: list[tuple[int, int, bool]] = []

    def callback(current: int, total: int, _src: Path, success: bool, _msg: str) -> None:
        progress.append((current, total, success))

    result = convert_batch([SAMPLE_DOCX], out_dir, callback=callback)

    assert result["success_count"] == 1
    assert progress == [(1, 1, True)]
    assert (out_dir / "sample.md").exists()


def test_convert_empty_result_raises(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    """Raise ConversionError when docx conversion returns empty content."""

    def fake_convert_docx(_src: Path, _dst: Path) -> str:
        raise ConversionError("Conversion produced empty output")

    monkeypatch.setattr("tomarkdown.converter._convert_docx", fake_convert_docx)

    with pytest.raises(ConversionError, match="empty output"):
        convert_docx_to_md(SAMPLE_DOCX, tmp_path / "out.md")


def test_convert_docx_extracts_images(tmp_path: Path) -> None:
    """Embedded images are written to an assets folder with relative links."""
    problem_docx = Path(
        r"E:\W-公司资料\项目分析\2026-05-激光散热系统\激光散热系统需求-芯蚁修订_20260521.docx"
    )
    if not problem_docx.exists():
        pytest.skip("sample problem docx not available on this machine")

    dst = tmp_path / "laser.md"
    convert_docx_to_md(problem_docx, dst)

    content = dst.read_text(encoding="utf-8")
    assets_dir = tmp_path / "laser_assets"

    assert "激光散热系统需求" in content
    assert assets_dir.is_dir()
    assert len(list(assets_dir.glob("image*.*"))) >= 1
    assert f"{assets_dir.name}/image1." in content
    assert "base64..." not in content
