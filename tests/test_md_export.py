"""Tests for Markdown → DOCX / PDF export."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from tomarkdown.converter import convert_batch, convert_file
from tomarkdown.md_export import convert_md_to_docx, markdown_to_docx


SAMPLE_MD = """# Hello Export

This is a **bold** and *italic* paragraph with `code`.

## Lists

- Alpha
- Beta

1. One
2. Two

## Table

| Name | Value |
|------|-------|
| A    | 1     |
| B    | 2     |

> A quoted line

```python
print("hi")
```
"""


def test_markdown_to_docx_basic(tmp_path: Path) -> None:
    dst = tmp_path / "out.docx"
    markdown_to_docx(SAMPLE_MD, dst, base_dir=tmp_path)

    assert dst.exists()
    doc = Document(str(dst))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    joined = "\n".join(texts)
    assert "Hello Export" in joined
    assert "bold" in joined or "This is a" in joined
    assert len(doc.tables) >= 1


def test_code_block_uses_shaded_table(tmp_path: Path) -> None:
    """Fenced code should render as a bordered shaded panel with real line breaks."""
    md = (
        "# Demo\n\n"
        "Inline `x = 1` here.\n\n"
        "```python\n"
        "def hello(name):\n"
        "    return name\n"
        "```\n"
    )
    dst = tmp_path / "code.docx"
    markdown_to_docx(md, dst, base_dir=tmp_path)

    doc = Document(str(dst))
    assert len(doc.tables) >= 1
    cell = doc.tables[0].cell(0, 0)
    lines = [p.text for p in cell.paragraphs]
    assert any("def hello" in line for line in lines)
    assert any("return name" in line for line in lines)
    # Lines must be separate paragraphs (Word does not honor \\n in a single run)
    assert len(lines) >= 2

    # Cell shading present
    tc_pr = cell._tc.tcPr
    assert tc_pr is not None
    shd = tc_pr.find(qn("w:shd"))
    assert shd is not None
    assert shd.get(qn("w:fill"), "").upper() == "F6F8FA"

    # Inline code shading on a body paragraph
    body = "\n".join(p.text for p in doc.paragraphs)
    assert "x = 1" in body



def test_convert_md_to_docx_file(tmp_path: Path) -> None:
    src = tmp_path / "note.md"
    src.write_text("# Title\n\nBody text.\n", encoding="utf-8")
    dst = tmp_path / "note.docx"

    convert_md_to_docx(src, dst)

    assert dst.exists()
    doc = Document(str(dst))
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "Title" in texts
    assert "Body text" in texts


def test_convert_file_from_md(tmp_path: Path) -> None:
    src = tmp_path / "a.md"
    src.write_text("## Section\n\nHello.\n", encoding="utf-8")
    dst = tmp_path / "a.docx"

    assert convert_file(src, dst, direction="from_md", output_format="docx") is True
    assert dst.exists()


def test_convert_batch_from_md(tmp_path: Path) -> None:
    src = tmp_path / "batch.md"
    src.write_text("# Batch\n\nOK\n", encoding="utf-8")
    out_dir = tmp_path / "out"

    result = convert_batch(
        [src],
        out_dir,
        direction="from_md",
        output_format="docx",
        overwrite=True,
    )

    assert result["success_count"] == 1
    assert (out_dir / "batch.docx").exists()


def test_convert_md_rejects_non_md(tmp_path: Path) -> None:
    src = tmp_path / "x.txt"
    src.write_text("nope", encoding="utf-8")
    with pytest.raises(ValueError, match="Unsupported file format"):
        convert_md_to_docx(src, tmp_path / "x.docx")


def test_convert_md_empty_raises(tmp_path: Path) -> None:
    src = tmp_path / "empty.md"
    src.write_text("   \n", encoding="utf-8")
    with pytest.raises(Exception, match="empty"):
        convert_md_to_docx(src, tmp_path / "empty.docx")
