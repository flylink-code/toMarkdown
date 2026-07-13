"""Tests for export settings persistence and header/footer application."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from tomarkdown.export_settings import (
    ExportSettings,
    load_export_settings,
    normalize_hex_color,
    save_export_settings,
)
from tomarkdown.md_export import markdown_to_docx


def test_save_and_load_export_settings(tmp_path: Path) -> None:
    path = tmp_path / "export_settings.json"
    settings = ExportSettings(
        output_format="pdf",
        header_enabled=True,
        header_text="公司文档 · {filename}",
        footer_text="第 {page} 页",
        doc_author="Tester",
    )
    save_export_settings(settings, path)

    loaded = load_export_settings(path)
    assert loaded.output_format == "pdf"
    assert loaded.format_label() == "PDF"
    assert loaded.header_enabled is True
    assert loaded.header_text == "公司文档 · {filename}"
    assert loaded.footer_text == "第 {page} 页"
    assert loaded.doc_author == "Tester"


def test_load_defaults_output_format(tmp_path: Path) -> None:
    path = tmp_path / "export_settings.json"
    path.write_text('{"header_enabled": true}', encoding="utf-8")
    loaded = load_export_settings(path)
    assert loaded.output_format == "docx"
    assert loaded.format_label() == "Word (.docx)"


def test_style_preset_and_heading_color(tmp_path: Path) -> None:
    settings = ExportSettings()
    settings.apply_style_preset("classic")
    assert settings.style_preset == "classic"
    assert settings.code_highlight is False
    assert settings.heading_h1_color == "000000"

    md = "# Red Title\n\n```python\nprint(1)\n```\n"
    dst = tmp_path / "styled.docx"
    settings.heading_h1_color = "CC0000"
    settings.code_bg = "EEF2FF"
    settings.code_show_language = False
    settings.style_preset = "custom"
    markdown_to_docx(md, dst, base_dir=tmp_path, settings=settings)

    doc = Document(str(dst))
    heading = next(p for p in doc.paragraphs if p.text.strip() == "Red Title")
    assert any(r.font.color.rgb and str(r.font.color.rgb) == "CC0000" for r in heading.runs)

    cell = doc.tables[0].cell(0, 0)
    shd = cell._tc.tcPr.find(qn("w:shd"))
    assert shd is not None
    assert shd.get(qn("w:fill"), "").upper() == "EEF2FF"
    # language label disabled → no "python" paragraph immediately before table body alone
    body_texts = [p.text.strip().lower() for p in doc.paragraphs if p.text.strip()]
    assert "python" not in body_texts


def test_normalize_hex_color() -> None:
    assert normalize_hex_color("#abc", "000000") == "AABBCC"
    assert normalize_hex_color("1f2328", "000000") == "1F2328"
    assert normalize_hex_color("xyz", "112233") == "112233"
    assert normalize_hex_color("12345", "112233") == "112233"


def test_markdown_docx_applies_header_footer(tmp_path: Path) -> None:
    src = tmp_path / "report.md"
    dst = tmp_path / "report.docx"
    settings = ExportSettings(
        header_enabled=True,
        header_text="页眉-{filename}",
        header_align="left",
        footer_enabled=True,
        footer_text="作者:{author} 第 {page} 页",
        footer_align="center",
        doc_author="Alice",
        doc_title="自定义标题",
    )

    markdown_to_docx(
        "# Hello\n\nBody\n",
        dst,
        base_dir=tmp_path,
        source=src,
        settings=settings,
    )

    doc = Document(str(dst))
    assert doc.core_properties.title == "自定义标题"
    assert doc.core_properties.author == "Alice"

    header_text = "".join(p.text for p in doc.sections[0].header.paragraphs)
    assert "页眉-report" in header_text

    footer_text = "".join(p.text for p in doc.sections[0].footer.paragraphs)
    assert "Alice" in footer_text
