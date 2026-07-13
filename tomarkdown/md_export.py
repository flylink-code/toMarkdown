"""Markdown → DOCX / PDF export (no GUI dependencies)."""

from __future__ import annotations

import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from markdown_it import MarkdownIt

from tomarkdown.converter import ConversionError, _find_libreoffice
from tomarkdown.export_settings import ExportSettings, build_placeholder_context

_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}

_PLACEHOLDER_RE = re.compile(
    r"\{(filename|filename_ext|date|title|author|page|numpages)\}"
)


_CODE_FONT = "Consolas"
_CODE_BG = "F6F8FA"
_CODE_BORDER = "D0D7DE"
_INLINE_CODE_BG = "F3F4F6"
_INLINE_CODE_FG = RGBColor(0xCF, 0x22, 0x2E)

# Pygments token colors inspired by GitHub Light / Typora default export
_TOKEN_COLORS: dict[str, RGBColor] = {
    "Token.Keyword": RGBColor(0xCF, 0x22, 0x2E),
    "Token.Keyword.Constant": RGBColor(0x05, 0x50, 0xAE),
    "Token.Keyword.Namespace": RGBColor(0xCF, 0x22, 0x2E),
    "Token.Keyword.Type": RGBColor(0x95, 0x38, 0x00),
    "Token.Name.Class": RGBColor(0x95, 0x38, 0x00),
    "Token.Name.Function": RGBColor(0x82, 0x50, 0xDF),
    "Token.Name.Builtin": RGBColor(0x05, 0x50, 0xAE),
    "Token.Name.Decorator": RGBColor(0x05, 0x50, 0xAE),
    "Token.String": RGBColor(0x0A, 0x30, 0x69),
    "Token.String.Affix": RGBColor(0x0A, 0x30, 0x69),
    "Token.String.Escape": RGBColor(0x0A, 0x30, 0x69),
    "Token.Comment": RGBColor(0x6E, 0x77, 0x81),
    "Token.Comment.Single": RGBColor(0x6E, 0x77, 0x81),
    "Token.Comment.Multiline": RGBColor(0x6E, 0x77, 0x81),
    "Token.Number": RGBColor(0x05, 0x50, 0xAE),
    "Token.Operator": RGBColor(0xCF, 0x22, 0x2E),
    "Token.Punctuation": RGBColor(0x24, 0x29, 0x2F),
    "Token.Name.Tag": RGBColor(0x11, 0x6B, 0x32),
    "Token.Name.Attribute": RGBColor(0x05, 0x50, 0xAE),
    "Token.Literal.String.Doc": RGBColor(0x0A, 0x30, 0x69),
}

_DEFAULT_CODE_FG = RGBColor(0x24, 0x29, 0x2F)


def _hex_to_rgb(value: str, fallback: RGBColor = _DEFAULT_CODE_FG) -> RGBColor:
    raw = (value or "").strip().lstrip("#")
    if len(raw) != 6:
        return fallback
    try:
        return RGBColor(int(raw[0:2], 16), int(raw[2:4], 16), int(raw[4:6], 16))
    except ValueError:
        return fallback


def _style_defaults(settings: ExportSettings | None) -> ExportSettings:
    if settings is None:
        settings = ExportSettings()
    settings.normalize_style_fields()
    return settings


def _set_run_font(run, *, east_asia: str = "微软雅黑", ascii_font: str = "Calibri") -> None:
    """Set Western + East-Asian fonts on a run."""
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia)


def _set_run_shading(run, fill: str) -> None:
    """Apply character background shading (Typora-like inline code chip)."""
    rpr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    rpr.append(shd)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_borders(cell, color: str = _CODE_BORDER) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
        borders.append(element)
    tc_pr.append(borders)


def _set_cell_margins(cell, *, top: int = 60, bottom: int = 60, left: int = 100, right: int = 100) -> None:
    """Set cell margins in twips (dxa)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for name, value in (
        ("top", top),
        ("left", left),
        ("bottom", bottom),
        ("right", right),
    ):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tc_pr.append(margins)


def _style_code_font(run, settings: ExportSettings) -> None:
    font = settings.code_font or _CODE_FONT
    size_pt = settings.code_font_size or 9.5
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size_pt)


def _token_color(token_type) -> RGBColor:
    """Resolve Pygments token type to a color by walking the type hierarchy."""
    current = token_type
    while current is not None:
        key = str(current)
        if key in _TOKEN_COLORS:
            return _TOKEN_COLORS[key]
        current = current.parent
    return _DEFAULT_CODE_FG


def _get_lexer(info: str):
    """Return a Pygments lexer for fence info string, or None."""
    lang = (info or "").strip().split()[0] if (info or "").strip() else ""
    if not lang:
        return None
    try:
        from pygments.lexers import get_lexer_by_name

        return get_lexer_by_name(lang, stripall=False)
    except Exception:
        return None


def _add_code_runs(paragraph, text: str, lexer, settings: ExportSettings) -> None:
    """Add monospace runs, optionally syntax-highlighted (single visual line)."""
    if not text:
        run = paragraph.add_run(" ")
        _style_code_font(run, settings)
        run.font.color.rgb = _DEFAULT_CODE_FG
        return

    if lexer is None:
        run = paragraph.add_run(text)
        _style_code_font(run, settings)
        run.font.color.rgb = _DEFAULT_CODE_FG
        return

    from pygments import lex

    # Caller owns line breaks — strip any newlines Pygments may attach to tokens.
    for token_type, value in lex(text, lexer):
        value = value.replace("\r", "").replace("\n", "")
        if not value:
            continue
        run = paragraph.add_run(value)
        _style_code_font(run, settings)
        run.font.color.rgb = _token_color(token_type)
        if "Comment" in str(token_type):
            run.italic = True


def _apply_inline(
    paragraph,
    tokens: list,
    base_dir: Path,
    settings: ExportSettings,
) -> None:
    """Render inline markdown-it tokens into a paragraph."""
    strong = 0
    em = 0
    code = 0
    link_href: str | None = None

    for token in tokens:
        ttype = token.type

        if ttype == "text":
            run = paragraph.add_run(token.content)
            _style_run(run, settings, strong=strong, em=em, code=code)
            if link_href:
                run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
                run.underline = True
        elif ttype == "code_inline":
            # Typora-like padded chip: thin spaces around content + shading
            run = paragraph.add_run(f"\u2009{token.content}\u2009")
            _style_run(run, settings, strong=strong, em=em, code=1)
        elif ttype == "softbreak":
            paragraph.add_run(" ")
        elif ttype == "hardbreak":
            paragraph.add_run().add_break(WD_BREAK.LINE)
        elif ttype == "strong_open":
            strong += 1
        elif ttype == "strong_close":
            strong = max(0, strong - 1)
        elif ttype == "em_open":
            em += 1
        elif ttype == "em_close":
            em = max(0, em - 1)
        elif ttype == "link_open":
            link_href = token.attrs.get("href") if token.attrs else None
        elif ttype == "link_close":
            if link_href:
                run = paragraph.add_run(f" ({link_href})")
                _set_run_font(run)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            link_href = None
        elif ttype == "image":
            _add_image(paragraph, token, base_dir)
        elif ttype == "html_inline":
            continue
        else:
            if token.children:
                _apply_inline(paragraph, token.children, base_dir, settings)


def _style_run(
    run,
    settings: ExportSettings,
    *,
    strong: int = 0,
    em: int = 0,
    code: int = 0,
) -> None:
    if code:
        _style_code_font(run, settings)
        run.font.color.rgb = _hex_to_rgb(settings.inline_code_fg, _INLINE_CODE_FG)
        _set_run_shading(run, settings.inline_code_bg or _INLINE_CODE_BG)
    else:
        _set_run_font(run)
    run.bold = strong > 0
    run.italic = em > 0


def _add_image(paragraph, token, base_dir: Path) -> None:
    src = (token.attrs or {}).get("src", "")
    alt = token.content or (token.attrs or {}).get("alt", "") or "image"
    if not src:
        paragraph.add_run(f"[{alt}]")
        return

    path = Path(src)
    if not path.is_absolute():
        path = (base_dir / path).resolve()

    if path.is_file():
        try:
            run = paragraph.add_run()
            run.add_picture(str(path), width=Inches(5.5))
            return
        except Exception:
            pass

    run = paragraph.add_run(f"![{alt}]({src})")
    _set_run_font(run)


def _add_code_block(
    doc: Document,
    content: str,
    info: str = "",
    *,
    settings: ExportSettings,
) -> None:
    """
    Render a fenced code block like Typora/GitHub DOCX export:
    rounded-feel gray panel (single-cell table), monospace lines, optional highlight.
    """
    language = (info or "").strip().split()[0] if (info or "").strip() else ""
    lines = content.replace("\r\n", "\n").replace("\r", "\n").rstrip("\n").split("\n")
    if not lines:
        lines = [""]

    before = doc.add_paragraph()
    before.paragraph_format.space_before = Pt(8)
    before.paragraph_format.space_after = Pt(0)
    before_run = before.add_run("")
    _set_run_font(before_run)

    if language and settings.code_show_language:
        label = doc.add_paragraph()
        label.paragraph_format.space_before = Pt(0)
        label.paragraph_format.space_after = Pt(2)
        run = label.add_run(language.lower())
        _set_run_font(run, ascii_font="Calibri")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x65, 0x6D, 0x76)
        run.italic = True

    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, settings.code_bg or _CODE_BG)
    _set_cell_borders(cell, settings.code_border or _CODE_BORDER)
    _set_cell_margins(cell, top=80, bottom=80, left=120, right=120)

    lexer = _get_lexer(language) if settings.code_highlight else None

    first = True
    for line in lines:
        if first:
            paragraph = cell.paragraphs[0]
            paragraph.clear()
            first = False
        else:
            paragraph = cell.add_paragraph()

        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.2
        paragraph.paragraph_format.left_indent = Pt(0)
        _add_code_runs(paragraph, line, lexer, settings)

    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(6)


def _add_heading(
    doc: Document,
    level: int,
    children: list,
    base_dir: Path,
    settings: ExportSettings,
) -> None:
    level = max(1, min(level, 9))
    paragraph = doc.add_heading(level=min(level, 9))
    if paragraph.runs:
        for run in paragraph.runs:
            run.text = ""
    _apply_inline(paragraph, children, base_dir, settings)
    color = _hex_to_rgb(settings.heading_color_for_level(level))
    for run in paragraph.runs:
        if not run.font.name:
            _set_run_font(run)
        run.font.color.rgb = color


def _extract_first_h1(md_text: str) -> str:
    """Return first level-1 heading text from markdown, if any."""
    for line in md_text.splitlines():
        stripped = line.strip()
        if stripped.startswith("# ") and not stripped.startswith("##"):
            return stripped[2:].strip()
    return ""


def _add_field_run(paragraph, instruction: str) -> None:
    """Append a Word field (e.g. PAGE / NUMPAGES) to paragraph."""
    run = paragraph.add_run()
    _set_run_font(run, ascii_font="Calibri")
    run.font.size = Pt(9)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    # Placeholder visible before field update
    text_elem = OxmlElement("w:t")
    text_elem.text = "1" if instruction == "PAGE" else ""

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text_elem)
    run._r.append(fld_end)


def _fill_header_footer_paragraph(
    paragraph,
    template: str,
    context: dict[str, str],
    *,
    align: str,
) -> None:
    """Render template with static placeholders and PAGE/NUMPAGES fields."""
    paragraph.clear()
    paragraph.alignment = _ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.CENTER)

    pos = 0
    for match in _PLACEHOLDER_RE.finditer(template):
        if match.start() > pos:
            run = paragraph.add_run(template[pos : match.start()])
            _set_run_font(run)
            run.font.size = Pt(9)
        key = match.group(1)
        if key == "page":
            _add_field_run(paragraph, "PAGE")
        elif key == "numpages":
            _add_field_run(paragraph, "NUMPAGES")
        else:
            run = paragraph.add_run(context.get(key, ""))
            _set_run_font(run)
            run.font.size = Pt(9)
        pos = match.end()

    if pos < len(template):
        run = paragraph.add_run(template[pos:])
        _set_run_font(run)
        run.font.size = Pt(9)


def _apply_export_settings(
    doc: Document,
    settings: ExportSettings,
    *,
    source: Path | None,
    inferred_title: str,
) -> None:
    """Apply document properties and header/footer chrome."""
    title = (settings.doc_title or inferred_title or (source.stem if source else "")).strip()
    author = settings.doc_author.strip()
    subject = settings.doc_subject.strip()

    core = doc.core_properties
    if title:
        core.title = title
    if author:
        core.author = author
    if subject:
        core.subject = subject

    context = build_placeholder_context(source=source, title=title, author=author)
    section = doc.sections[0]
    section.different_first_page_header_footer = settings.different_first_page

    if settings.header_enabled and settings.header_text.strip():
        header = section.header
        header.is_linked_to_previous = False
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        _fill_header_footer_paragraph(
            paragraph,
            settings.header_text,
            context,
            align=settings.header_align,
        )
    if settings.footer_enabled and settings.footer_text.strip():
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        _fill_header_footer_paragraph(
            paragraph,
            settings.footer_text,
            context,
            align=settings.footer_align,
        )

    if settings.different_first_page:
        if settings.first_header_text.strip() or settings.header_enabled:
            first_header = section.first_page_header
            first_header.is_linked_to_previous = False
            paragraph = (
                first_header.paragraphs[0]
                if first_header.paragraphs
                else first_header.add_paragraph()
            )
            text = settings.first_header_text if settings.first_header_text.strip() else ""
            if text:
                _fill_header_footer_paragraph(
                    paragraph,
                    text,
                    context,
                    align=settings.header_align,
                )
            else:
                paragraph.clear()
        if settings.first_footer_text.strip() or settings.footer_enabled:
            first_footer = section.first_page_footer
            first_footer.is_linked_to_previous = False
            paragraph = (
                first_footer.paragraphs[0]
                if first_footer.paragraphs
                else first_footer.add_paragraph()
            )
            text = settings.first_footer_text if settings.first_footer_text.strip() else ""
            if text:
                _fill_header_footer_paragraph(
                    paragraph,
                    text,
                    context,
                    align=settings.footer_align,
                )
            else:
                paragraph.clear()


def markdown_to_docx(
    md_text: str,
    dst: Path,
    *,
    base_dir: Path | None = None,
    source: Path | None = None,
    settings: ExportSettings | None = None,
) -> None:
    """Convert markdown text to a .docx file."""
    base_dir = Path(base_dir) if base_dir is not None else Path.cwd()
    dst = Path(dst)
    settings = _style_defaults(settings)

    md = MarkdownIt("commonmark", {"breaks": True, "html": True}).enable("table")
    tokens = md.parse(md_text)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "微软雅黑")

    i = 0
    while i < len(tokens):
        token = tokens[i]
        ttype = token.type

        if ttype == "heading_open":
            level = int(token.tag[1]) if token.tag and token.tag.startswith("h") else 1
            inline = tokens[i + 1] if i + 1 < len(tokens) else None
            children = inline.children if inline and inline.type == "inline" else []
            _add_heading(doc, level, children or [], base_dir, settings)
            i += 3  # heading_open, inline, heading_close
            continue

        if ttype == "paragraph_open":
            inline = tokens[i + 1] if i + 1 < len(tokens) else None
            paragraph = doc.add_paragraph()
            if inline and inline.type == "inline" and inline.children:
                _apply_inline(paragraph, inline.children, base_dir, settings)
            i += 3
            continue

        if ttype == "fence" or ttype == "code_block":
            _add_code_block(
                doc,
                token.content,
                getattr(token, "info", "") or "",
                settings=settings,
            )
            i += 1
            continue

        if ttype == "hr":
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        if ttype == "blockquote_open":
            i += 1
            while i < len(tokens) and tokens[i].type != "blockquote_close":
                if tokens[i].type == "paragraph_open":
                    inline = tokens[i + 1] if i + 1 < len(tokens) else None
                    paragraph = doc.add_paragraph()
                    paragraph.paragraph_format.left_indent = Inches(0.3)
                    if inline and inline.type == "inline" and inline.children:
                        _apply_inline(paragraph, inline.children, base_dir, settings)
                    for run in paragraph.runs:
                        run.italic = True
                        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                    i += 3
                else:
                    i += 1
            i += 1  # blockquote_close
            continue

        if ttype in {"bullet_list_open", "ordered_list_open"}:
            ordered = ttype == "ordered_list_open"
            i += 1
            item_index = 1
            while i < len(tokens) and tokens[i].type not in {
                "bullet_list_close",
                "ordered_list_close",
            }:
                if tokens[i].type == "list_item_open":
                    i += 1
                    while i < len(tokens) and tokens[i].type != "list_item_close":
                        if tokens[i].type == "paragraph_open":
                            inline = tokens[i + 1] if i + 1 < len(tokens) else None
                            style_name = "List Number" if ordered else "List Bullet"
                            try:
                                paragraph = doc.add_paragraph(style=style_name)
                            except KeyError:
                                paragraph = doc.add_paragraph()
                                prefix = f"{item_index}. " if ordered else "• "
                                run = paragraph.add_run(prefix)
                                _set_run_font(run)
                            if inline and inline.type == "inline" and inline.children:
                                _apply_inline(
                                    paragraph, inline.children, base_dir, settings
                                )
                            item_index += 1
                            i += 3
                        elif tokens[i].type in {"bullet_list_open", "ordered_list_open"}:
                            # Nested lists: flatten one level with indent.
                            nested_ordered = tokens[i].type == "ordered_list_open"
                            i += 1
                            nested_idx = 1
                            while i < len(tokens) and tokens[i].type not in {
                                "bullet_list_close",
                                "ordered_list_close",
                            }:
                                if tokens[i].type == "list_item_open":
                                    i += 1
                                    while i < len(tokens) and tokens[i].type != "list_item_close":
                                        if tokens[i].type == "paragraph_open":
                                            inline = (
                                                tokens[i + 1] if i + 1 < len(tokens) else None
                                            )
                                            paragraph = doc.add_paragraph()
                                            paragraph.paragraph_format.left_indent = Inches(0.5)
                                            prefix = (
                                                f"{nested_idx}. " if nested_ordered else "◦ "
                                            )
                                            run = paragraph.add_run(prefix)
                                            _set_run_font(run)
                                            if (
                                                inline
                                                and inline.type == "inline"
                                                and inline.children
                                            ):
                                                _apply_inline(
                                                    paragraph,
                                                    inline.children,
                                                    base_dir,
                                                    settings,
                                                )
                                            nested_idx += 1
                                            i += 3
                                        else:
                                            i += 1
                                    i += 1
                                else:
                                    i += 1
                            i += 1
                        else:
                            i += 1
                    i += 1  # list_item_close
                else:
                    i += 1
            i += 1  # list close
            continue

        if ttype == "table_open":
            i = _render_table(doc, tokens, i, base_dir, settings)
            continue

        if ttype == "html_block":
            i += 1
            continue

        i += 1

    _apply_export_settings(
        doc,
        settings,
        source=source,
        inferred_title=_extract_first_h1(md_text),
    )

    dst.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(dst))
    except OSError as exc:
        raise ConversionError(f"Failed to write output file {dst}: {exc}") from exc


def _render_table(
    doc: Document,
    tokens: list,
    start: int,
    base_dir: Path,
    settings: ExportSettings,
) -> int:
    """Render a markdown table starting at table_open; return next index."""
    i = start + 1
    rows: list[list[list]] = []
    current_row: list[list] | None = None

    while i < len(tokens) and tokens[i].type != "table_close":
        ttype = tokens[i].type
        if ttype == "tr_open":
            current_row = []
            i += 1
        elif ttype == "tr_close":
            if current_row is not None:
                rows.append(current_row)
            current_row = None
            i += 1
        elif ttype in {"th_open", "td_open"}:
            inline = tokens[i + 1] if i + 1 < len(tokens) else None
            children = inline.children if inline and inline.type == "inline" else []
            if current_row is not None:
                current_row.append(children or [])
            i += 3  # open, inline, close
        else:
            i += 1

    if rows:
        cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=cols)
        table.style = "Table Grid"
        for r_idx, row_cells in enumerate(rows):
            for c_idx in range(cols):
                cell = table.rows[r_idx].cells[c_idx]
                cell.text = ""
                paragraph = cell.paragraphs[0]
                children = row_cells[c_idx] if c_idx < len(row_cells) else []
                if children:
                    _apply_inline(paragraph, children, base_dir, settings)
                if r_idx == 0:
                    for run in paragraph.runs:
                        run.bold = True

    return i + 1  # skip table_close


def convert_md_to_docx(
    src: Path,
    dst: Path,
    *,
    settings: ExportSettings | None = None,
) -> None:
    """Convert a .md file to .docx."""
    src = Path(src)
    dst = Path(dst)
    if not src.exists():
        raise FileNotFoundError(f"Source file not found: {src}")
    if not src.is_file():
        raise ValueError(f"Source is not a file: {src}")
    if src.suffix.lower() not in {".md", ".markdown"}:
        raise ValueError(f"Unsupported file format (expected .md): {src.suffix}")

    try:
        text = src.read_text(encoding="utf-8")
    except OSError as exc:
        raise ConversionError(f"Failed to read {src}: {exc}") from exc

    if not text.strip():
        raise ConversionError("Conversion produced empty output")

    markdown_to_docx(
        text,
        dst,
        base_dir=src.parent,
        source=src,
        settings=settings,
    )


def _docx_to_pdf_via_word(src: Path, dst: Path) -> None:
    """Convert .docx to .pdf using Microsoft Word COM automation."""
    import pythoncom
    import win32com.client

    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(src.resolve()))
        # 17 = wdFormatPDF
        doc.SaveAs2(str(dst.resolve()), FileFormat=17)
        doc.Close()
    finally:
        if doc is not None:
            doc = None
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()


def _docx_to_pdf_via_libreoffice(src: Path, out_dir: Path) -> Path:
    """Convert .docx to .pdf using LibreOffice headless mode."""
    soffice = _find_libreoffice()
    if soffice is None:
        raise ConversionError("LibreOffice not found")

    result = subprocess.run(
        [
            str(soffice),
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(out_dir),
            str(src.resolve()),
        ],
        capture_output=True,
        text=True,
        check=False,
    )
    if result.returncode != 0:
        stderr = result.stderr.strip() or result.stdout.strip() or "unknown error"
        raise ConversionError(f"LibreOffice conversion failed: {stderr}")

    dst = out_dir / f"{src.stem}.pdf"
    if not dst.exists():
        raise ConversionError("LibreOffice did not produce a .pdf file")
    return dst


def convert_md_to_pdf(
    src: Path,
    dst: Path,
    *,
    settings: ExportSettings | None = None,
) -> None:
    """
    Convert a .md file to .pdf via intermediate DOCX, then Word or LibreOffice.
    """
    src = Path(src)
    dst = Path(dst)
    if not src.exists():
        raise FileNotFoundError(f"Source file not found: {src}")
    if src.suffix.lower() not in {".md", ".markdown"}:
        raise ValueError(f"Unsupported file format (expected .md): {src.suffix}")

    temp_dir = Path(tempfile.mkdtemp(prefix="tomarkdown_mdpdf_"))
    errors: list[str] = []
    try:
        docx_path = temp_dir / f"{src.stem}.docx"
        convert_md_to_docx(src, docx_path, settings=settings)

        if sys.platform == "win32":
            try:
                dst.parent.mkdir(parents=True, exist_ok=True)
                _docx_to_pdf_via_word(docx_path, dst)
                if dst.exists():
                    return
            except Exception as exc:
                errors.append(f"Word: {exc}")

        try:
            produced = _docx_to_pdf_via_libreoffice(docx_path, temp_dir)
            dst.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(produced, dst)
            return
        except Exception as exc:
            errors.append(f"LibreOffice: {exc}")

        raise ConversionError(
            "无法将 Markdown 转为 PDF。请安装 Microsoft Word 或 LibreOffice。"
            + (f" ({'; '.join(errors)})" if errors else "")
        )
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)
